"""Document ingestion pipeline: parse -> chunk -> embed -> index in Redis.

Supported formats: PDF (via ``pypdf``) and plain text. Each chunk is stored as
a Redis HASH so RediSearch can run KNN vector queries over the embeddings.
"""

from __future__ import annotations

import io
import uuid
from datetime import datetime, timezone

import numpy as np
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

from app.config import PUBLIC_USER_ID, get_settings
from app.services.embeddings import get_embeddings
from app.services.redis_client import get_redis


class UnsupportedFileType(ValueError):
    """Raised when an uploaded file has an unsupported extension."""


class TextExtractionError(ValueError):
    """Raised when no text could be extracted from a document."""


class EmbeddingError(RuntimeError):
    """Raised when the embedding provider returns malformed vectors."""


def _extract_pdf_text(data: bytes) -> str:
    """Extract the embedded text layer from a PDF (empty for scanned PDFs)."""
    reader = PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(data: bytes) -> str:
    """Extract text from a DOCX file, including paragraphs and table cells."""
    document = DocxDocument(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _ocr_available() -> bool:
    """True if the optional OCR stack (PyMuPDF + pytesseract) is installed."""
    try:
        import fitz  # noqa: F401
        import pytesseract  # noqa: F401
    except ImportError:
        return False
    return True


def _ocr_pdf(data: bytes) -> str:  # pragma: no cover - requires native OCR stack
    """Render each PDF page to an image and OCR it with Tesseract."""
    import fitz
    import pytesseract
    from PIL import Image

    settings = get_settings()
    document = fitz.open(stream=data, filetype="pdf")
    pages: list[str] = []
    for page in document:
        pixmap = page.get_pixmap(dpi=settings.ocr_dpi)
        image = Image.open(io.BytesIO(pixmap.tobytes("png")))
        pages.append(pytesseract.image_to_string(image, lang=settings.ocr_language))
    return "\n".join(pages)


def _extract_text(filename: str, data: bytes) -> str:
    """Extract raw text from a PDF or TXT payload.

    For PDFs without a text layer (e.g. scanned documents) it falls back to OCR
    when the optional OCR stack is installed; otherwise it raises a clear error.
    """
    lower = filename.lower()
    if lower.endswith(".pdf"):
        text = _extract_pdf_text(data)
        if text.strip():
            return text
        if _ocr_available():
            return _ocr_pdf(data)
        raise TextExtractionError(
            "Não foi possível extrair texto do PDF (provavelmente escaneado). "
            "Habilite o OCR reconstruindo a imagem com INSTALL_OCR=true."
        )
    if lower.endswith(".docx"):
        return _extract_docx_text(data)
    if lower.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")
    raise UnsupportedFileType(f"Unsupported file type: {filename}")


def chunk_text(text: str) -> list[str]:
    """Split text into overlapping chunks sized by tokens.

    ``CHUNK_SIZE``/``CHUNK_OVERLAP`` are measured in tokens (tiktoken) rather
    than characters, matching how embedding/LLM context budgets are counted.
    """
    settings = get_settings()
    splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
    )
    return [c for c in splitter.split_text(text) if c.strip()]


def _to_float32_bytes(vector: list[float]) -> bytes:
    """Serialize an embedding vector to the float32 bytes Redis expects."""
    return np.asarray(vector, dtype=np.float32).tobytes()


def _validate_embeddings(vectors: list[list[float]]) -> None:
    """Reject empty or ragged embedding output before indexing.

    A provider returning vectors of inconsistent length (or empty vectors)
    would corrupt the index and make KNN search fail silently, so we surface it
    as a clear error instead.
    """
    if not vectors or not vectors[0]:
        raise EmbeddingError("Embedding provider returned no vectors.")
    dim = len(vectors[0])
    if any(len(v) != dim for v in vectors):
        raise EmbeddingError("Embedding provider returned vectors of varying length.")


def prepare_file(
    filename: str,
    data: bytes,
    api_key: str | None = None,
) -> dict:
    """Extract, chunk and embed a file **without** writing to Redis.

    Returns a dict consumed by :func:`index_prepared`. Any parsing/validation
    error (``UnsupportedFileType``, ``TextExtractionError``, ``EmbeddingError``)
    is raised here, before indexing, so a multi-file upload can be made
    all-or-nothing: nothing is persisted unless every file prepares cleanly.
    """
    text = _extract_text(filename, data)
    chunks = chunk_text(text)
    prepared: dict = {
        "file_id": str(uuid.uuid4()),
        "name": filename,
        "chunks": chunks,
        "vectors": [],
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
    }
    if not chunks:
        return prepared

    vectors = get_embeddings(api_key).embed_documents(chunks)
    _validate_embeddings(vectors)
    prepared["vectors"] = vectors
    return prepared


def index_prepared(prepared: dict, user_id: str = PUBLIC_USER_ID) -> dict:
    """Write an already-prepared file's chunks to Redis (owner-tagged).

    Returns a dict with ``file_id``, ``name``, ``chunks_indexed`` and ``status``
    so the upload endpoint can serialize it directly.
    """
    chunks = prepared["chunks"]
    if not chunks:
        return {
            "file_id": prepared["file_id"],
            "name": prepared["name"],
            "chunks_indexed": 0,
            "status": "empty",
        }

    settings = get_settings()
    file_id = prepared["file_id"]
    client = get_redis()
    pipe = client.pipeline()

    for idx, (chunk, vector) in enumerate(zip(chunks, prepared["vectors"])):
        key = f"{settings.redis_key_prefix}{file_id}:chunk:{idx}"
        pipe.hset(
            key,
            mapping={
                "content": chunk,
                "source": prepared["name"],
                "file_id": file_id,
                "user_id": user_id,
                "chunk_index": idx,
                "uploaded_at": prepared["uploaded_at"],
                "embedding": _to_float32_bytes(vector),
            },
        )
    pipe.execute()

    return {
        "file_id": file_id,
        "name": prepared["name"],
        "chunks_indexed": len(chunks),
        "status": "indexed",
    }


def ingest_file(
    filename: str,
    data: bytes,
    api_key: str | None = None,
    user_id: str = PUBLIC_USER_ID,
) -> dict:
    """Run the full ingestion pipeline for a single file (prepare + index).

    ``api_key`` optionally overrides the embedding provider key per request.
    ``user_id`` tags every chunk so the document is only visible to its owner
    (``PUBLIC_USER_ID`` when auth is disabled).
    """
    return index_prepared(prepare_file(filename, data, api_key), user_id)
