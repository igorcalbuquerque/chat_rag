"""Tests for the ingestion service: chunking, embeddings and Redis indexing."""

from __future__ import annotations

import pytest

from app.config import get_settings
from app.services import ingestion


def test_chunk_text_respects_size_and_overlap(monkeypatch):
    import tiktoken

    monkeypatch.setenv("CHUNK_SIZE", "50")
    monkeypatch.setenv("CHUNK_OVERLAP", "10")
    get_settings.cache_clear()

    text = "palavra " * 200
    chunks = ingestion.chunk_text(text)

    # Chunking is token-based: each chunk must stay within the token budget.
    enc = tiktoken.get_encoding("cl100k_base")
    assert len(chunks) > 1
    assert all(len(enc.encode(c)) <= 50 for c in chunks)


def test_chunk_text_skips_empty():
    assert ingestion.chunk_text("   \n  ") == []


def test_extract_text_txt():
    text = ingestion._extract_text("notes.txt", b"hello world")
    assert text == "hello world"


def test_extract_text_unsupported():
    with pytest.raises(ingestion.UnsupportedFileType):
        ingestion._extract_text("image.png", b"\x00\x01")


def _blank_pdf_bytes() -> bytes:
    import io

    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def test_extract_text_pdf_with_text_layer(monkeypatch):
    monkeypatch.setattr(ingestion, "_extract_pdf_text", lambda data: "conteudo real")
    assert ingestion._extract_text("doc.pdf", b"%PDF-fake") == "conteudo real"


def test_extract_text_pdf_without_text_raises_when_no_ocr(monkeypatch):
    monkeypatch.setattr(ingestion, "_ocr_available", lambda: False)
    with pytest.raises(ingestion.TextExtractionError):
        ingestion._extract_text("scan.pdf", _blank_pdf_bytes())


def test_extract_text_pdf_falls_back_to_ocr(monkeypatch):
    monkeypatch.setattr(ingestion, "_ocr_available", lambda: True)
    monkeypatch.setattr(ingestion, "_ocr_pdf", lambda data: "texto via ocr")
    assert ingestion._extract_text("scan.pdf", _blank_pdf_bytes()) == "texto via ocr"


def test_ocr_available_true(monkeypatch):
    import sys
    import types

    monkeypatch.setitem(sys.modules, "fitz", types.ModuleType("fitz"))
    monkeypatch.setitem(sys.modules, "pytesseract", types.ModuleType("pytesseract"))
    assert ingestion._ocr_available() is True


def test_ocr_available_false(monkeypatch):
    import sys

    monkeypatch.setitem(sys.modules, "fitz", None)  # import fitz -> ImportError
    assert ingestion._ocr_available() is False


def _docx_bytes(paragraphs, table_rows=None) -> bytes:
    import io

    from docx import Document

    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, value in enumerate(row):
                table.cell(r, c).text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_extract_text_docx():
    data = _docx_bytes(
        ["Primeiro parágrafo.", "Segundo parágrafo."],
        table_rows=[["A1", "B1"]],
    )
    text = ingestion._extract_text("doc.docx", data)
    assert "Primeiro parágrafo." in text
    assert "Segundo parágrafo." in text
    assert "A1" in text and "B1" in text


def test_ingest_docx_indexes_chunks(fake_redis, fake_embeddings):
    data = _docx_bytes(["conteudo de teste " * 50])
    result = ingestion.ingest_file("relatorio.docx", data)
    assert result["status"] == "indexed"
    assert result["chunks_indexed"] >= 1


def test_ingest_file_indexes_chunks(fake_redis, fake_embeddings):
    data = ("frase de teste. " * 200).encode("utf-8")
    result = ingestion.ingest_file("doc.txt", data)

    assert result["status"] == "indexed"
    assert result["chunks_indexed"] > 0
    assert result["name"] == "doc.txt"

    # Each chunk is stored as a hash with the expected fields.
    prefix = get_settings().redis_key_prefix
    keys = list(fake_redis.scan_iter(match=f"{prefix}*"))
    assert len(keys) == result["chunks_indexed"]

    fields = fake_redis.hgetall(keys[0])
    assert b"content" in fields
    assert b"embedding" in fields
    assert fields[b"source"] == b"doc.txt"


def test_ingest_empty_file(fake_redis, fake_embeddings):
    result = ingestion.ingest_file("empty.txt", b"   ")
    assert result["chunks_indexed"] == 0
    assert result["status"] == "empty"
